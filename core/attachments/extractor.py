"""Safe, bounded text extraction for user-provided attachments."""

from __future__ import annotations

import base64
import hashlib
import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from common.contracts import Attachment, AttachmentExtraction, SourceDocument

_MAX_ATTACHMENT_BYTES = 10 * 1024 * 1024
_MAX_EXTRACTED_CHARACTERS = 100_000
_MAX_CONTEXT_CHARACTERS = 30_000
_TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".html", ".htm"}


def _load_bytes(attachment: Attachment) -> bytes:
    if attachment.content_base64:
        return base64.b64decode(attachment.content_base64, validate=True)
    if attachment.source_path:
        path = Path(attachment.source_path).expanduser().resolve()
        if not path.is_file():
            raise FileNotFoundError(f"attachment path does not exist: {path}")
        if path.stat().st_size > _MAX_ATTACHMENT_BYTES:
            raise ValueError("attachment exceeds 10 MB extraction limit")
        return path.read_bytes()
    raise ValueError("attachment has no readable content")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        raise ValueError("encrypted PDF is not supported")
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                chunks.append(" | ".join(values))
    return "\n".join(chunks).strip()


def _extract_text(data: bytes, filename: str, content_type: str | None) -> str:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf" or content_type == "application/pdf":
        return _extract_pdf(data)
    if extension == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(data)
    if extension in _TEXT_EXTENSIONS or (content_type or "").startswith("text/"):
        return data.decode("utf-8-sig", errors="replace").strip()
    raise NotImplementedError(f"unsupported attachment type: {extension or content_type or 'unknown'}")


def extract_attachments(attachments: list[Attachment]) -> tuple[list[SourceDocument], list[AttachmentExtraction]]:
    documents: list[SourceDocument] = []
    results: list[AttachmentExtraction] = []
    for attachment in attachments:
        try:
            data = _load_bytes(attachment)
            if len(data) > _MAX_ATTACHMENT_BYTES:
                raise ValueError("attachment exceeds 10 MB extraction limit")
            text = _extract_text(data, attachment.filename, attachment.content_type)
            if not text:
                results.append(AttachmentExtraction(
                    attachment_id=attachment.attachment_id,
                    filename=attachment.filename,
                    status="empty",
                ))
                continue
            truncated = len(text) > _MAX_EXTRACTED_CHARACTERS
            text = text[:_MAX_EXTRACTED_CHARACTERS]
            digest = hashlib.sha1(f"{attachment.attachment_id}:{attachment.filename}".encode("utf-8")).hexdigest()[:16]
            documents.append(SourceDocument(
                doc_id=f"attachment:{digest}",
                source_id=f"attachment:{attachment.filename}",
                title=attachment.filename,
                url=f"attachment://{attachment.attachment_id}",
                content=text,
            ))
            results.append(AttachmentExtraction(
                attachment_id=attachment.attachment_id,
                filename=attachment.filename,
                status="extracted",
                character_count=len(text),
                truncated=truncated,
            ))
        except NotImplementedError as exc:
            results.append(AttachmentExtraction(
                attachment_id=attachment.attachment_id,
                filename=attachment.filename,
                status="unsupported",
                error=str(exc),
            ))
        except Exception as exc:  # noqa: BLE001 - one bad attachment must not discard the others
            results.append(AttachmentExtraction(
                attachment_id=attachment.attachment_id,
                filename=attachment.filename,
                status="failed",
                error=str(exc),
            ))
    return documents, results


def build_question_context(question: str, documents: list[SourceDocument]) -> str:
    if not documents:
        return question
    chunks = [question, "\n\n[첨부자료 본문]"]
    remaining = _MAX_CONTEXT_CHARACTERS
    for document in documents:
        content = document.content or ""
        excerpt = content[:remaining]
        chunks.append(f"\n--- {document.title} ({document.doc_id}) ---\n{excerpt}")
        remaining -= len(excerpt)
        if remaining <= 0:
            break
    return "".join(chunks)
